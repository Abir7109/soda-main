import os, csv, json, io, asyncio
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

EXPORT_DIR = Path.home() / "Desktop" / "soda_exports"

_SODA_BG = RGBColor(0x0A, 0x0E, 0x17)
_SODA_CYAN = RGBColor(0x00, 0xFB, 0xFB)
_SODA_TEXT = RGBColor(0xCC, 0xCC, 0xCC)
_SODA_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_SODA_LINE = RGBColor(0x1A, 0x2A, 0x3A)

def _ensure_export_dir():
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)

def _data_to_rows(data):
    if isinstance(data, str):
        try:
            data = json.loads(data)
        except json.JSONDecodeError:
            return [{"content": data}]
    if isinstance(data, dict):
        if all(isinstance(v, (list, dict, str, int, float, bool)) for v in data.values()):
            return [data]
        for k, v in data.items():
            if isinstance(v, list) and len(v) > 0 and isinstance(v[0], dict):
                return v
        return [data]
    if isinstance(data, list):
        if len(data) == 0:
            return [{"content": "No data"}]
        if all(isinstance(x, dict) for x in data):
            return data
        return [{"item": str(x)} for x in data]
    return [{"data": str(data)}]

def _write_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)

def _set_cell_text(cell, text, font_name="Consolas", font_size=9, bold=False, color=_SODA_TEXT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color
    _remove_space_before(p)
    _remove_space_after(p)

def _remove_space_before(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), "0")

def _remove_space_after(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:after"), "40")

def _set_page_background(doc, color_hex):
    sectPr = doc.sections[0]._sectPr
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), color_hex)
    sectPr.append(bg)

def _add_heading_custom(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sizes = {1: 18, 2: 14, 3: 12}
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(sizes.get(level, 12))
    run.font.bold = True
    run.font.color.rgb = _SODA_CYAN
    _remove_space_after(p)
    return p

def _add_line(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "00fbfb")
    pBdr.append(bottom)
    pPr.append(pBdr)
    _remove_space_before(p)
    _remove_space_after(p)

def _add_body_text(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(str(text))
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = _SODA_TEXT
    _remove_space_after(p)

def export_markdown(data, title):
    _ensure_export_dir()
    rows = _data_to_rows(data)
    lines = [f"# {title}", "", "---", ""]
    if len(rows) == 1 and list(rows[0].keys()) == ["content"]:
        lines.append(rows[0]["content"])
        lines.append("")
    else:
        headers = list(rows[0].keys())
        lines.append("| " + " | ".join(h.replace("_", " ").title() for h in headers) + " |")
        lines.append("| " + " | ".join("---" for _ in headers) + " |")
        for row in rows:
            vals = [str(row.get(h, "")) for h in headers]
            lines.append("| " + " | ".join(v.replace("\n", " ")[:80] for v in vals) + " |")
        lines.append("")
        lines.append(f"*{len(rows)} records*")
        lines.append("")
    path = EXPORT_DIR / f"{title}.md"
    path.write_text("\n".join(lines), encoding="utf-8")
    return {"success": True, "path": str(path), "message": f"Saved markdown with {len(rows)} records"}

def export_csv(data, title):
    _ensure_export_dir()
    rows = _data_to_rows(data)
    path = EXPORT_DIR / f"{title}.csv"
    if not rows:
        path.write_text("", encoding="utf-8")
        return {"success": True, "path": str(path), "message": "Empty CSV"}
    headers = list(rows[0].keys())
    with open(path, "w", newline="", encoding="utf-8-sig") as f:
        w = csv.DictWriter(f, fieldnames=headers)
        w.writeheader()
        w.writerows({h: row.get(h, "") for h in headers} for row in rows)
    return {"success": True, "path": str(path), "message": f"Saved CSV with {len(rows)} records"}

def export_docx(data, title):
    _ensure_export_dir()
    rows = _data_to_rows(data)
    path = EXPORT_DIR / f"{title}.docx"
    doc = Document()
    _set_page_background(doc, "0A0E17")
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
    _add_heading_custom(doc, title, 1)
    _add_line(doc)
    if len(rows) == 1 and list(rows[0].keys()) == ["content"]:
        _add_body_text(doc, rows[0]["content"])
    else:
        headers = list(rows[0].keys())
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        hdr = table.rows[0]
        for i, h in enumerate(headers):
            cell = hdr.cells[i]
            _write_cell_shading(cell, "001a2e")
            _set_cell_text(cell, h.replace("_", " ").title(), bold=True, color=_SODA_CYAN, font_size=9)
            width = Inches(9.0 / len(headers))
            cell.width = width
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, h in enumerate(headers):
                cell = row_cells[i]
                val = str(row_data.get(h, ""))
                _set_cell_text(cell, val, font_size=8)
        _add_line(doc)
        p = doc.add_paragraph()
        run = p.add_run(f"{len(rows)} records  |  SODA Export")
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = _SODA_LINE
    doc.save(str(path))
    return {"success": True, "path": str(path), "message": f"Saved DOCX with {len(rows)} records"}

async def export_data(data, export_format, title):
    if not title:
        title = "soda_export"
    title = "".join(c if c.isalnum() or c in " _-" else "_" for c in title)[:60]
    export_format = export_format.lower().strip()
    if export_format in ("md", "markdown", "mdown"):
        return export_markdown(data, title)
    elif export_format in ("csv",):
        return export_csv(data, title)
    elif export_format in ("docx", "doc", "word"):
        result = await asyncio.to_thread(export_docx, data, title)
        return result
    else:
        return {"success": False, "message": f"Unsupported format: {export_format}. Use markdown, csv, or docx."}
